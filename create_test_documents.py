from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

# Create a new document
doc = Document()

# Set landscape orientation
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE

# Add content to make it clear this is a test document
doc.add_heading('Landscape Document - Rotated 90 Degrees', 0)
doc.add_paragraph('This is a test document created to test rotation detection.')
doc.add_paragraph('The document should be in landscape orientation but will appear rotated 90 degrees when opened.')
doc.add_paragraph('')

# Add some sample content
doc.add_paragraph('Sample content line 1')
doc.add_paragraph('Sample content line 2')
doc.add_paragraph('Sample content line 3')

# Add a table to make the landscape orientation more obvious
table = doc.add_table(rows=3, cols=4)
for i in range(3):
    for j in range(4):
        cell = table.cell(i, j)
        cell.text = f'Row {i+1}, Col {j+1}'

doc.add_paragraph('')
doc.add_paragraph('End of test document')

# Save the document
doc.save('landscape_test_document.docx')
print("Created landscape_test_document.docx")