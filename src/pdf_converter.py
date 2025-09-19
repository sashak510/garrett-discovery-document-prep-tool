"""
PDF Converter Module
Converts various document formats to PDF with OCR capabilities
"""

import os
import sys
from pathlib import Path
import logging
import tempfile
import shutil
from datetime import datetime

try:
    from PIL import Image
except ImportError:
    Image = None

import pytesseract
import imutils
import numpy as np

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    canvas = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


class PDFConverter:
    """Converts various document formats to PDF with OCR support and enhanced line detection"""
    
    def __init__(self, log_callback=None, tesseract_path=None):
        """
        Initialize the PDF converter
        
        Args:
            log_callback: Optional callback function for logging messages
            tesseract_path: Path to tesseract executable (for OCR)
        """
        self.log_callback = log_callback
        self.conversion_errors = []
        
        # Document type classification for processing strategy
        self.document_types = {
            'HIGH_ACCURACY': ['word', 'text'],  # 100% accurate line detection
            'READABLE_PDF': ['pdf_text'],       # High accuracy with content analysis
            'COMPLEX_PDF': ['pdf_image', 'pdf_ocr'],  # Needs fallback strategies
            'SCANNED': ['tiff', 'image']        # OCR + fallback required
        }
        
        # Set up tesseract path if provided
        if tesseract_path and pytesseract:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            
        # Check for required dependencies
        self._check_dependencies()
        
    def log(self, message):
        """Log a message using the callback or print"""
        if self.log_callback:
            self.log_callback(message)
        else:
            print(message)
            
    def _check_dependencies(self):
        """Check if required dependencies are available"""
        missing_deps = []
        
        if not Image:
            missing_deps.append("Pillow (PIL)")
        if not pytesseract:
            missing_deps.append("pytesseract")
        if not Document:
            missing_deps.append("python-docx")
        if not canvas:
            missing_deps.append("reportlab")
        if not fitz:
            missing_deps.append("PyMuPDF")
            
        if missing_deps:
            self.log(f"Warning: Missing dependencies: {', '.join(missing_deps)}")
            self.log("Some conversion features may not work properly")
            
        # Try to find tesseract
        if pytesseract:
            try:
                pytesseract.get_tesseract_version()
                self.log("Tesseract OCR found and working")
            except Exception as e:
                self.log(f"Warning: Tesseract OCR not found or not working: {e}")
                self.log("OCR functionality will not be available")
                
    def classify_document_type(self, input_path):
        """
        Classify document type for optimal processing strategy
        
        Args:
            input_path (str): Path to input file
            
        Returns:
            tuple: (document_category, document_subtype, confidence_score)
        """
        input_file = Path(input_path)
        file_ext = input_file.suffix.lower()
        
        try:
            if file_ext in ['.docx', '.doc']:
                return 'HIGH_ACCURACY', 'word', 1.0
            elif file_ext == '.txt':
                return 'HIGH_ACCURACY', 'text', 1.0
            elif file_ext == '.pdf':
                # Analyze PDF to determine if it's text-based or image-based
                pdf_type, confidence, warnings = self._analyze_pdf_content(input_path)
                if pdf_type == 'text_based':
                    return 'READABLE_PDF', 'pdf_text', confidence
                elif pdf_type == 'mixed':
                    return 'COMPLEX_PDF', 'pdf_mixed', confidence
                else:
                    return 'COMPLEX_PDF', 'pdf_image', confidence
            elif file_ext in ['.tiff', '.tif']:
                return 'SCANNED', 'tiff', 1.0
            else:
                return 'COMPLEX_PDF', 'unknown', 0.5
                
        except Exception as e:
            self.log(f"Error classifying document type: {e}")
            return 'COMPLEX_PDF', 'unknown', 0.0
    
    def _analyze_pdf_content(self, pdf_path):
        """
        Analyze PDF to determine content type and detect unusual layouts
        
        Args:
            pdf_path (str): Path to PDF file
            
        Returns:
            tuple: (content_type, confidence_score, layout_warnings)
        """
        if not fitz:
            return 'image_based', 0.5, []
            
        try:
            doc = fitz.open(pdf_path)
            total_chars = 0
            total_images = 0
            total_pages = doc.page_count
            layout_warnings = []
            
            # Sample first 5 pages for analysis
            sample_pages = min(5, total_pages)
            
            # Track layout characteristics across pages
            rotated_text_pages = 0
            multi_column_pages = 0
            empty_with_numbers_pages = 0
            
            for page_num in range(sample_pages):
                page = doc[page_num]
                
                # Count extractable text characters
                text = page.get_text()
                total_chars += len(text.strip())
                
                # Count images
                images = page.get_images()
                total_images += len(images)
                
                # DETECT UNUSUAL LAYOUTS
                page_warnings = self._detect_unusual_layout(page, page_num + 1)
                
                # Track specific layout issues
                for warning in page_warnings:
                    if "rotated" in warning.lower():
                        rotated_text_pages += 1
                    elif "column" in warning.lower():
                        multi_column_pages += 1
                    elif "empty" in warning.lower() and "numbers" in warning.lower():
                        empty_with_numbers_pages += 1
                
                layout_warnings.extend(page_warnings)
            
            doc.close()
            
            # Add summary warnings for recurring issues
            if rotated_text_pages > 0:
                layout_warnings.append(f"⚠️  {rotated_text_pages}/{sample_pages} pages have rotated text")
            if multi_column_pages > 1:
                layout_warnings.append(f"⚠️  {multi_column_pages}/{sample_pages} pages have multi-column layout")
            if empty_with_numbers_pages > 0:
                layout_warnings.append(f"⚠️  {empty_with_numbers_pages}/{sample_pages} pages appear empty but may have line numbers")
            
            # Classification logic
            chars_per_page = total_chars / sample_pages
            images_per_page = total_images / sample_pages
            
            if chars_per_page > 200:  # Rich text content
                if images_per_page < 2:
                    return 'text_based', 0.9, layout_warnings
                else:
                    return 'mixed', 0.7, layout_warnings
            elif chars_per_page > 50:  # Some text content
                # If there are images present, it's likely a scanned document with OCR text
                if images_per_page > 0:
                    return 'image_based', 0.8, layout_warnings
                else:
                    return 'mixed', 0.6, layout_warnings
            else:  # Minimal or no text
                return 'image_based', 0.8, layout_warnings
                
        except Exception as e:
            self.log(f"Error analyzing PDF content: {e}")
            return 'image_based', 0.3, [f"❌ Analysis failed: {str(e)}"]

    def _detect_unusual_layout(self, page, page_num):
        """
        Detect unusual PDF layouts that may cause line numbering issues
        
        Args:
            page: PyMuPDF page object
            page_num: Page number for reporting
            
        Returns:
            list: List of warning messages about unusual layouts
        """
        warnings = []
        
        try:
            # Get text blocks with position information
            text_dict = page.get_text("dict")
            blocks = text_dict.get("blocks", [])
            
            if not blocks:
                warnings.append(f"📄 Page {page_num}: No text blocks found (may be image-only)")
                return warnings
            
            # Analyze text blocks for unusual patterns
            text_blocks = [b for b in blocks if "lines" in b]
            
            # 1. DETECT ROTATED TEXT (90-degree rotation)
            rotated_spans = 0
            total_spans = 0
            
            # 2. DETECT MULTI-COLUMN LAYOUT
            x_positions = []
            
            # 3. DETECT EMPTY PAGES WITH EXISTING LINE NUMBERS
            digit_only_text = []
            
            for block in text_blocks:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        total_spans += 1
                        
                        # Check text content
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                            
                        # Get span bbox for position analysis
                        bbox = span.get("bbox", [0, 0, 0, 0])
                        x_positions.append(bbox[0])  # Left edge position
                        
                        # ROTATED TEXT DETECTION
                        # Check if text orientation is unusual
                        flags = span.get("flags", 0)
                        if flags & 2**3:  # Text rotation flag (rough heuristic)
                            rotated_spans += 1
                        
                        # Alternative: Check character spacing/width ratios
                        # Only flag as potentially rotated if characters are extremely tall (likely actually rotated)
                        char_width = (bbox[2] - bbox[0]) / max(len(text), 1)
                        char_height = bbox[3] - bbox[1]
                        if char_height > char_width * 4:  # Much higher threshold for actual rotation
                            rotated_spans += 1
                        
                        # EXISTING LINE NUMBERS DETECTION
                        # Check for standalone numbers (possible existing line numbers)
                        if text.isdigit() and len(text) <= 4:
                            digit_only_text.append(int(text))
            
            # ANALYSIS & WARNINGS
            
            # 1. Rotated text warning
            if total_spans > 0:
                rotation_ratio = rotated_spans / total_spans
                if rotation_ratio > 0.3:  # More than 30% of text appears rotated
                    warnings.append(f"🔄 Page {page_num}: {rotation_ratio:.1%} of text appears rotated 90°")
            
            # 2. Multi-column detection
            if len(x_positions) > 10:  # Need sufficient data points
                # Find distinct column positions (cluster X coordinates)
                x_positions.sort()
                columns = []
                current_column = x_positions[0]
                
                for x in x_positions[1:]:
                    if x - current_column > 100:  # New column if >100pt gap
                        columns.append(current_column)
                        current_column = x
                columns.append(current_column)
                
                if len(columns) >= 2:
                    warnings.append(f"📊 Page {page_num}: Multi-column layout detected ({len(columns)} columns)")
            
            # 3. Empty page with existing line numbers
            non_digit_text = [text for block in text_blocks 
                            for line in block.get("lines", []) 
                            for span in line.get("spans", [])
                            if (text := span.get("text", "").strip()) and not text.isdigit()]
            
            if len(digit_only_text) > 5 and len(non_digit_text) < 3:
                # Check if digits form a sequence (likely existing line numbers)
                digit_only_text.sort()
                if len(digit_only_text) > 1:
                    gaps = [digit_only_text[i+1] - digit_only_text[i] for i in range(len(digit_only_text)-1)]
                    avg_gap = sum(gaps) / len(gaps) if gaps else 0
                    
                    if 0.5 <= avg_gap <= 2:  # Sequential-ish numbers
                        warnings.append(f"📝 Page {page_num}: Appears empty but contains {len(digit_only_text)} sequential numbers (existing line numbers?)")
            
            # 4. UNUSUAL ASPECT RATIOS OR PAGE SIZES
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height
            aspect_ratio = page_width / page_height if page_height > 0 else 1
            
            if aspect_ratio > 1.5 or aspect_ratio < 0.5:
                warnings.append(f"📐 Page {page_num}: Unusual aspect ratio {aspect_ratio:.2f} (may affect line numbering)")
            
        except Exception as e:
            warnings.append(f"❌ Page {page_num}: Layout analysis failed - {str(e)}")
        
        return warnings

    def convert_to_pdf(self, input_path, output_path, perform_ocr=True):
        """
        Convert a file to PDF format using optimal strategy based on document type
        
        Args:
            input_path (str): Path to input file
            output_path (str): Path for output PDF file
            perform_ocr (bool): Whether to perform OCR on images
            
        Returns:
            tuple: (success: bool, document_type: str, processing_notes: str)
        """
        input_file = Path(input_path)
        output_file = Path(output_path)
        
        if not input_file.exists():
            self.log(f"Error: Input file does not exist: {input_path}")
            return False, 'unknown', 'File not found'
            
        # Ensure output directory exists
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Classify document type for optimal processing
        doc_category, doc_subtype, confidence = self.classify_document_type(input_path)
        self.log(f"Document classified as {doc_category}/{doc_subtype} (confidence: {confidence:.2f})")
        
        file_ext = input_file.suffix.lower()
        processing_notes = f"Processed as {doc_category}/{doc_subtype}"
        
        try:
            if file_ext == '.pdf':
                success = self._handle_existing_pdf(input_path, output_path, perform_ocr, doc_subtype)
            elif file_ext in ['.tiff', '.tif']:
                success = self._convert_tiff_to_pdf(input_path, output_path, perform_ocr)
            elif file_ext in ['.docx', '.doc']:
                success = self._convert_word_to_pdf_enhanced(input_path, output_path)
                processing_notes += " - Enhanced line mapping"
            elif file_ext == '.txt':
                success = self._convert_text_to_pdf_enhanced(input_path, output_path)
                processing_notes += " - Enhanced line mapping"
            elif file_ext == '.rtf':
                success = self._convert_rtf_to_pdf(input_path, output_path)
            else:
                self.log(f"Error: Unsupported file format: {file_ext}")
                return False, doc_subtype, 'Unsupported format'
                
            return success, doc_subtype, processing_notes
                
        except Exception as e:
            self.log(f"Error converting {input_path}: {str(e)}")
            self.conversion_errors.append({
                'file': input_path,
                'error': str(e),
                'type': 'conversion_error'
            })
            return False, doc_subtype, f'Conversion error: {str(e)}'
            
    def _handle_existing_pdf(self, input_path, output_path, perform_ocr=True, doc_subtype='pdf_text'):
        """Handle PDF files with strategy based on document type"""
        # Check if input and output are the same file
        input_resolved = Path(input_path).resolve()
        output_resolved = Path(output_path).resolve()
        
        if input_resolved == output_resolved:
            self.log(f"PDF already in correct location (no conversion needed): {Path(input_path).name}")
            return True
            
        if not perform_ocr:
            # Just copy the file
            shutil.copy2(input_path, output_path)
            self.log(f"Copied PDF: {Path(input_path).name}")
            return True
        
        # Processing strategy based on document type
        if doc_subtype == 'pdf_text':
            # High-quality text PDF - minimal processing needed
            shutil.copy2(input_path, output_path)
            self.log(f"Text-based PDF copied (high line accuracy expected): {Path(input_path).name}")
            return True
        elif doc_subtype == 'pdf_mixed':
            # Mixed content - copy but note potential line accuracy issues
            shutil.copy2(input_path, output_path)
            self.log(f"Mixed-content PDF copied (moderate line accuracy): {Path(input_path).name}")
            return True
        else:
            # Image-based or unknown - may need OCR
            if self._pdf_needs_ocr(input_path):
                return self._ocr_pdf(input_path, output_path)
            else:
                shutil.copy2(input_path, output_path)
                self.log(f"PDF copied with fallback processing: {Path(input_path).name}")
                return True
            
    def _pdf_needs_ocr(self, pdf_path):
        """Check if a PDF needs OCR (has no extractable text)"""
        if not fitz:
            return False
            
        try:
            doc = fitz.open(pdf_path)
            total_text = ""
            
            # Check first few pages for text
            for page_num in range(min(3, doc.page_count)):
                page = doc[page_num]
                text = page.get_text()
                total_text += text.strip()
                
            doc.close()
            
            # If we found minimal text, it probably needs OCR
            return len(total_text) < 50
            
        except Exception as e:
            self.log(f"Error checking PDF text content: {e}")
            return False
            
    def _ocr_pdf(self, input_path, output_path):
        """Perform OCR on a PDF file"""
        if not pytesseract or not fitz:
            self.log("OCR not available - copying PDF as-is")
            shutil.copy2(input_path, output_path)
            return True
            
        try:
            # Open the PDF
            doc = fitz.open(input_path)
            new_doc = fitz.open()  # Create new document
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                
                # Convert page to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Save image temporarily
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                    tmp_img.write(img_data)
                    tmp_img_path = tmp_img.name
                    
                try:
                    # NEW: Detect and correct orientation before OCR
                    corrected_img_path, rotation_applied = self._detect_and_correct_orientation(tmp_img_path)
                    
                    # Perform OCR on corrected image
                    ocr_text = pytesseract.image_to_string(Image.open(corrected_img_path))
                    
                    # Create new page with correct dimensions (swap if rotated 90/270)
                    if rotation_applied in (90, 270):
                        # Swap width and height for 90/270 degree rotations
                        new_page = new_doc.new_page(width=page.rect.height, height=page.rect.width)
                        page_rect = fitz.Rect(0, 0, page.rect.height, page.rect.width)
                    else:
                        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
                        page_rect = page.rect
                    
                    # Insert the corrected image
                    new_page.insert_image(page_rect, filename=corrected_img_path)
                    
                    # Add invisible text overlay for searchability
                    if ocr_text.strip():
                        text_rect = page_rect
                        new_page.insert_textbox(text_rect, ocr_text, 
                                              fontsize=8, color=(1, 1, 1),  # White text (invisible)
                                              overlay=True)
                    
                finally:
                    # Clean up temporary images
                    if os.path.exists(tmp_img_path):
                        os.unlink(tmp_img_path)
                    if 'corrected_img_path' in locals() and corrected_img_path != tmp_img_path and os.path.exists(corrected_img_path):
                        os.unlink(corrected_img_path)
                    
            # Save the new PDF
            new_doc.save(output_path)
            new_doc.close()
            doc.close()
            
            self.log(f"OCR completed for PDF: {Path(input_path).name}")
            return True
            
        except Exception as e:
            self.log(f"Error performing OCR on PDF: {e}")
            # Fallback: copy original
            shutil.copy2(input_path, output_path)
            return True
            
    def _convert_tiff_to_pdf(self, input_path, output_path, perform_ocr=True):
        """Convert TIFF file to PDF with optional OCR"""
        if not Image:
            self.log("PIL not available for TIFF conversion")
            return False
            
        try:
            # Open TIFF image
            with Image.open(input_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                    
                if perform_ocr and pytesseract:
                    # Save image temporarily for orientation detection
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                        img.save(tmp_img.name, "PNG")
                        tmp_img_path = tmp_img.name
                    
                    try:
                        # NEW: Detect and correct orientation before OCR
                        corrected_img_path, _ = self._detect_and_correct_orientation(tmp_img_path)
                        
                        # Load corrected image and perform OCR
                        corrected_img = Image.open(corrected_img_path)
                        ocr_text = pytesseract.image_to_string(corrected_img)
                        
                        # Create PDF with corrected image and text
                        self._create_pdf_with_image_and_text(corrected_img, ocr_text, output_path)
                        
                    finally:
                        # Clean up temporary images
                        if os.path.exists(tmp_img_path):
                            os.unlink(tmp_img_path)
                        if 'corrected_img_path' in locals() and corrected_img_path != tmp_img_path and os.path.exists(corrected_img_path):
                            os.unlink(corrected_img_path)
                else:
                    # Just convert image to PDF
                    img.save(output_path, "PDF", resolution=300.0)
                    
            self.log(f"Converted TIFF to PDF: {Path(input_path).name}")
            return True
            
        except Exception as e:
            self.log(f"Error converting TIFF: {e}")
            return False
            
    def _create_pdf_with_image_and_text(self, image, text, output_path):
        """Create a PDF with an image and searchable text overlay"""
        if not canvas:
            # Fallback: just save image as PDF
            image.save(output_path, "PDF", resolution=300.0)
            return
            
        try:
            # Save image temporarily
            with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp_img:
                image.save(tmp_img.name, "JPEG", quality=95)
                tmp_img_path = tmp_img.name
                
            # Create PDF
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Add image
            img_width, img_height = image.size
            aspect_ratio = img_height / img_width
            
            # Scale image to fit page
            if img_width > img_height:
                # Landscape orientation
                img_pdf_width = width - 72  # 1 inch margin
                img_pdf_height = img_pdf_width * aspect_ratio
            else:
                # Portrait orientation
                img_pdf_height = height - 72  # 1 inch margin
                img_pdf_width = img_pdf_height / aspect_ratio
                
            # Center image
            x = (width - img_pdf_width) / 2
            y = (height - img_pdf_height) / 2
            
            c.drawImage(tmp_img_path, x, y, img_pdf_width, img_pdf_height)
            
            # Add invisible text for searchability
            if text.strip():
                c.setFillColorRGB(1, 1, 1)  # White (invisible)
                c.setFont("Helvetica", 8)
                
                # Split text into lines and add to PDF
                lines = text.split('\n')
                text_y = height - 50
                for line in lines[:50]:  # Limit to first 50 lines
                    if line.strip():
                        c.drawString(50, text_y, line.strip()[:100])  # Limit line length
                        text_y -= 12
                        if text_y < 50:
                            break
                            
            c.save()
            
            # Clean up
            os.unlink(tmp_img_path)
            
        except Exception as e:
            self.log(f"Error creating PDF with text overlay: {e}")
            # Fallback
            image.save(output_path, "PDF", resolution=300.0)
            
    def _detect_and_correct_orientation(self, image_path):
        """
        Detect orientation and rotate image if needed
        
        Args:
            image_path (str): Path to the image file
            
        Returns:
            tuple: (corrected_image_path, rotation_applied) 
                   rotation_applied is degrees rotated (0, 90, 180, 270)
        """
        try:
            # Detect orientation using Tesseract OSD
            osd_data = pytesseract.image_to_osd(
                image_path,
                config='--psm 0 -c min_characters_to_try=5',
                output_type=pytesseract.Output.DICT
            )
            
            rotation_needed = osd_data['rotate']
            confidence = osd_data['orientation_conf']
            orientation = osd_data['orientation']
            
            # Log detailed detection results
            self.log(f"🔍 Tesseract OSD results: rotation={rotation_needed}°, confidence={confidence:.1f}, orientation={orientation}°")
            
            # Lower confidence threshold and add more detailed logging
            if confidence > 5.0 and rotation_needed != 0:
                # Load and rotate image
                image = Image.open(image_path)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Convert to numpy array for imutils
                image_np = np.array(image)
                
                # Rotate with bounds preservation
                rotated = imutils.rotate_bound(image_np, angle=rotation_needed)
                
                # Save corrected image
                path_obj = Path(image_path)
                corrected_path = str(path_obj.parent / f"{path_obj.stem}_corrected{path_obj.suffix}")
                Image.fromarray(rotated).save(corrected_path)
                
                self.log(f"🔄 Corrected orientation: rotated {rotation_needed}° (confidence: {confidence:.1f}, detected orientation: {orientation}°)")
                return corrected_path, rotation_needed
            else:
                if rotation_needed == 0:
                    self.log(f"✅ Document orientation correct (confidence: {confidence:.1f})")
                else:
                    self.log(f"⚠️  Orientation detection confidence too low ({confidence:.1f}) - skipping rotation")
                return image_path, 0
                
        except Exception as e:
            self.log(f"⚠️  Orientation detection failed: {e} - using original image")
            return image_path, 0  # Fallback to original
            
    def _convert_word_to_pdf_enhanced(self, input_path, output_path):
        """Convert Word document to PDF with precise line mapping for 100% accuracy"""
        if not Document or not canvas:
            self.log("Required libraries not available for Word conversion")
            return False
            
        try:
            # Read Word document
            doc = Document(input_path)
            
            # Extract paragraphs with enhanced processing
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
                elif len(paragraphs) > 0:  # Add blank line markers
                    paragraphs.append("")  # Preserve blank lines
                    
            # Create PDF with precise line tracking
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Enhanced settings for consistent line numbering
            font_size = 10
            line_height = 12  # Consistent with line numbering system
            left_margin = 72  # 1 inch
            right_margin = 72  # 1 inch
            top_margin = 72   # 1 inch
            bottom_margin = 72  # 1 inch
            usable_width = width - left_margin - right_margin
            
            c.setFont("Helvetica", font_size)
            y_position = height - top_margin
            
            # Track every line position for perfect alignment
            line_positions = []
            total_lines = 0
            
            for paragraph_text in paragraphs:
                if not paragraph_text:  # Empty paragraph = blank line
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Helvetica", font_size)
                    continue
                    
                # Smart word wrapping with exact line tracking
                words = paragraph_text.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    
                    # Check if line fits
                    if c.stringWidth(test_line, "Helvetica", font_size) <= usable_width:
                        current_line = test_line
                    else:
                        # Current line is full, output it
                        if current_line:
                            c.drawString(left_margin, y_position, current_line)
                            line_positions.append(y_position)
                            y_position -= line_height
                            total_lines += 1
                            
                            # Check for new page
                            if y_position < bottom_margin:
                                c.showPage()
                                y_position = height - top_margin
                                c.setFont("Helvetica", font_size)
                                
                        current_line = word
                        
                # Output the final line of the paragraph
                if current_line:
                    c.drawString(left_margin, y_position, current_line)
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Helvetica", font_size)
                        
                # Add space between paragraphs
                y_position -= line_height * 0.5
                        
            c.save()
            
            # Store line mapping metadata for later use
            self._store_line_mapping(output_path, line_positions, total_lines)
            
            self.log(f"Enhanced Word conversion completed: {Path(input_path).name} ({total_lines} lines mapped)")
            return True
            
        except Exception as e:
            self.log(f"Error in enhanced Word conversion: {e}")
            # Fallback to original method
            return self._convert_word_to_pdf(input_path, output_path)
            
    def _convert_word_to_pdf(self, input_path, output_path):
        """Convert Word document to PDF (fallback method)"""
        if not Document or not canvas:
            self.log("Required libraries not available for Word conversion")
            return False
            
        try:
            # Read Word document
            doc = Document(input_path)
            
            # Extract text
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
                
            # Create PDF
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Add text to PDF
            y_position = height - 72  # Start 1 inch from top
            line_height = 14
            
            c.setFont("Helvetica", 10)
            
            for paragraph_text in full_text:
                if not paragraph_text.strip():
                    y_position -= line_height
                    continue
                    
                # Word wrap
                words = paragraph_text.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    if c.stringWidth(test_line) < width - 144:  # 2 inch margins
                        current_line = test_line
                    else:
                        if current_line:
                            c.drawString(72, y_position, current_line)
                            y_position -= line_height
                            
                        current_line = word
                        
                        # Check if we need a new page
                        if y_position < 72:
                            c.showPage()
                            y_position = height - 72
                            c.setFont("Helvetica", 10)
                            
                # Draw the last line
                if current_line:
                    c.drawString(72, y_position, current_line)
                    y_position -= line_height * 1.5  # Extra space between paragraphs
                    
                # Check if we need a new page
                if y_position < 72:
                    c.showPage()
                    y_position = height - 72
                    c.setFont("Helvetica", 10)
                    
            c.save()
            
            self.log(f"Converted Word document to PDF: {Path(input_path).name}")
            return True
            
        except Exception as e:
            self.log(f"Error converting Word document: {e}")
            return False
            
    def _convert_text_to_pdf_enhanced(self, input_path, output_path):
        """Convert text file to PDF with precise line mapping for 100% accuracy"""
        if not canvas:
            self.log("ReportLab not available for text conversion")
            return False
            
        try:
            # Read text file preserving exact line structure
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
                
            # Remove only trailing newlines but preserve internal structure
            lines = [line.rstrip('\n\r') for line in lines]
                
            # Create PDF with precise line tracking
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Enhanced settings for consistent line numbering
            font_size = 9
            line_height = 12  # Consistent with line numbering system
            left_margin = 72  # 1 inch
            bottom_margin = 72  # 1 inch
            top_margin = 72   # 1 inch
            max_chars_per_line = 80  # Standard text width
            
            c.setFont("Courier", font_size)  # Monospace for text files
            y_position = height - top_margin
            
            # Track every line position for perfect alignment
            line_positions = []
            total_lines = 0
            
            for original_line in lines:
                # Handle empty lines
                if not original_line.strip():
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Courier", font_size)
                    continue
                
                # Handle long lines by wrapping but tracking each visual line
                remaining_text = original_line
                while remaining_text:
                    # Extract chunk that fits on one line
                    if len(remaining_text) <= max_chars_per_line:
                        line_chunk = remaining_text
                        remaining_text = ""
                    else:
                        # Smart break at word boundary if possible
                        break_point = max_chars_per_line
                        if remaining_text[max_chars_per_line] != ' ':
                            # Look backwards for space
                            for i in range(max_chars_per_line - 1, max_chars_per_line - 20, -1):
                                if i > 0 and remaining_text[i] == ' ':
                                    break_point = i
                                    break
                        
                        line_chunk = remaining_text[:break_point]
                        remaining_text = remaining_text[break_point:].lstrip()
                    
                    # Draw this line chunk
                    c.drawString(left_margin, y_position, line_chunk)
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Courier", font_size)
                    
            c.save()
            
            # Store line mapping metadata for later use
            self._store_line_mapping(output_path, line_positions, total_lines)
            
            self.log(f"Enhanced text conversion completed: {Path(input_path).name} ({total_lines} lines mapped)")
            return True
            
        except Exception as e:
            self.log(f"Error in enhanced text conversion: {e}")
            # Fallback to original method
            return self._convert_text_to_pdf(input_path, output_path)

    def _convert_text_to_pdf(self, input_path, output_path):
        """Convert text file to PDF (fallback method)"""
        if not canvas:
            self.log("ReportLab not available for text conversion")
            return False
            
        try:
            # Read text file
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = f.read()
                
            # Create PDF
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Add text to PDF
            y_position = height - 72  # Start 1 inch from top
            line_height = 12
            
            c.setFont("Courier", 9)  # Monospace font for text files
            
            lines = text_content.split('\n')
            
            for line in lines:
                # Handle long lines
                while len(line) > 80:  # Wrap at 80 characters
                    c.drawString(72, y_position, line[:80])
                    line = line[80:]
                    y_position -= line_height
                    
                    if y_position < 72:
                        c.showPage()
                        y_position = height - 72
                        c.setFont("Courier", 9)
                        
                # Draw the line
                c.drawString(72, y_position, line)
                y_position -= line_height
                
                # Check if we need a new page
                if y_position < 72:
                    c.showPage()
                    y_position = height - 72
                    c.setFont("Courier", 9)
                    
            c.save()
            
            self.log(f"Converted text file to PDF: {Path(input_path).name}")
            return True
            
        except Exception as e:
            self.log(f"Error converting text file: {e}")
            return False
            
    def _convert_rtf_to_pdf(self, input_path, output_path):
        """Convert RTF file to PDF (basic implementation)"""
        # For now, treat RTF as text file
        # A more sophisticated implementation would parse RTF formatting
        return self._convert_text_to_pdf(input_path, output_path)
    
    def _convert_word_to_pdf_enhanced(self, input_path, output_path):
        """Convert Word document to PDF with precise line mapping for 100% accuracy"""
        if not Document or not canvas:
            self.log("Required libraries not available for Word conversion")
            return False
            
        try:
            # Read Word document
            doc = Document(input_path)
            
            # Extract paragraphs with enhanced processing
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
                elif len(paragraphs) > 0:  # Add blank line markers
                    paragraphs.append("")  # Preserve blank lines
                    
            # Create PDF with precise line tracking
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Enhanced settings for consistent line numbering
            font_size = 10
            line_height = 12  # Consistent with line numbering system
            left_margin = 72  # 1 inch
            right_margin = 72  # 1 inch
            top_margin = 72   # 1 inch
            bottom_margin = 72  # 1 inch
            usable_width = width - left_margin - right_margin
            
            c.setFont("Helvetica", font_size)
            y_position = height - top_margin
            
            # Track every line position for perfect alignment
            line_positions = []
            total_lines = 0
            
            for paragraph_text in paragraphs:
                if not paragraph_text:  # Empty paragraph = blank line
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Helvetica", font_size)
                    continue
                    
                # Smart word wrapping with exact line tracking
                words = paragraph_text.split()
                current_line = ""
                
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    
                    # Check if line fits
                    if c.stringWidth(test_line, "Helvetica", font_size) <= usable_width:
                        current_line = test_line
                    else:
                        # Current line is full, output it
                        if current_line:
                            c.drawString(left_margin, y_position, current_line)
                            line_positions.append(y_position)
                            y_position -= line_height
                            total_lines += 1
                            
                            # Check for new page
                            if y_position < bottom_margin:
                                c.showPage()
                                y_position = height - top_margin
                                c.setFont("Helvetica", font_size)
                                
                        current_line = word
                        
                # Output the final line of the paragraph
                if current_line:
                    c.drawString(left_margin, y_position, current_line)
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Helvetica", font_size)
                        
                # Add space between paragraphs
                y_position -= line_height * 0.5
                        
            c.save()
            
            # Store line mapping metadata for later use
            self._store_line_mapping(output_path, line_positions, total_lines)
            
            self.log(f"Enhanced Word conversion completed: {Path(input_path).name} ({total_lines} lines mapped)")
            return True
            
        except Exception as e:
            self.log(f"Error in enhanced Word conversion: {e}")
            # Fallback to original method
            return self._convert_word_to_pdf(input_path, output_path)
    
    def _convert_text_to_pdf_enhanced(self, input_path, output_path):
        """Convert text file to PDF with precise line mapping for 100% accuracy"""
        if not canvas:
            self.log("ReportLab not available for text conversion")
            return False
            
        try:
            # Read text file preserving exact line structure
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
                
            # Remove only trailing newlines but preserve internal structure
            lines = [line.rstrip('\n\r') for line in lines]
                
            # Create PDF with precise line tracking
            c = canvas.Canvas(str(output_path), pagesize=letter)
            width, height = letter
            
            # Enhanced settings for consistent line numbering
            font_size = 9
            line_height = 12  # Consistent with line numbering system
            left_margin = 72  # 1 inch
            bottom_margin = 72  # 1 inch
            top_margin = 72   # 1 inch
            max_chars_per_line = 80  # Standard text width
            
            c.setFont("Courier", font_size)  # Monospace for text files
            y_position = height - top_margin
            
            # Track every line position for perfect alignment
            line_positions = []
            total_lines = 0
            
            for original_line in lines:
                # Handle empty lines
                if not original_line.strip():
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Courier", font_size)
                    continue
                
                # Handle long lines by wrapping but tracking each visual line
                remaining_text = original_line
                while remaining_text:
                    # Extract chunk that fits on one line
                    if len(remaining_text) <= max_chars_per_line:
                        line_chunk = remaining_text
                        remaining_text = ""
                    else:
                        # Smart break at word boundary if possible
                        break_point = max_chars_per_line
                        if remaining_text[max_chars_per_line] != ' ':
                            # Look backwards for space
                            for i in range(max_chars_per_line - 1, max_chars_per_line - 20, -1):
                                if i > 0 and remaining_text[i] == ' ':
                                    break_point = i
                                    break
                        
                        line_chunk = remaining_text[:break_point]
                        remaining_text = remaining_text[break_point:].lstrip()
                    
                    # Draw this line chunk
                    c.drawString(left_margin, y_position, line_chunk)
                    line_positions.append(y_position)
                    y_position -= line_height
                    total_lines += 1
                    
                    # Check for new page
                    if y_position < bottom_margin:
                        c.showPage()
                        y_position = height - top_margin
                        c.setFont("Courier", font_size)
                    
            c.save()
            
            # Store line mapping metadata for later use
            self._store_line_mapping(output_path, line_positions, total_lines)
            
            self.log(f"Enhanced text conversion completed: {Path(input_path).name} ({total_lines} lines mapped)")
            return True
            
        except Exception as e:
            self.log(f"Error in enhanced text conversion: {e}")
            # Fallback to original method
            return self._convert_text_to_pdf(input_path, output_path)
        
    def _store_line_mapping(self, pdf_path, line_positions, total_lines):
        """Store line mapping metadata for enhanced line numbering accuracy"""
        try:
            import json
            
            # Create metadata file path
            pdf_path_obj = Path(pdf_path)
            metadata_path = pdf_path_obj.with_suffix('.linemap.json')
            
            # Store line mapping data
            mapping_data = {
                'pdf_file': pdf_path_obj.name,
                'total_lines': total_lines,
                'line_positions': line_positions,
                'conversion_type': 'enhanced',
                'line_height': 12,
                'font_size': 10,
                'created_timestamp': str(datetime.now())
            }
            
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(mapping_data, f, indent=2)
                
            self.log(f"Line mapping metadata saved: {metadata_path.name}")
            
        except Exception as e:
            self.log(f"Warning: Could not save line mapping metadata: {e}")
            # Not critical - continue without metadata
            
    def load_line_mapping(self, pdf_path):
        """Load line mapping metadata if available"""
        try:
            import json
            
            pdf_path_obj = Path(pdf_path)
            metadata_path = pdf_path_obj.with_suffix('.linemap.json')
            
            if metadata_path.exists():
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    mapping_data = json.load(f)
                    
                self.log(f"Line mapping metadata loaded: {mapping_data['total_lines']} lines")
                return mapping_data
            else:
                return None
                
        except Exception as e:
            self.log(f"Warning: Could not load line mapping metadata: {e}")
            return None

    def get_conversion_errors(self):
        """Get list of conversion errors"""
        return self.conversion_errors
        
    def clear_errors(self):
        """Clear the conversion errors list"""
        self.conversion_errors = []


